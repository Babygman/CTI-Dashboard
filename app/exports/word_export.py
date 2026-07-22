from datetime import datetime
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

from flask import current_app
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.models.system_setting import SystemSetting


class WordExporter:
    """Generate enterprise Word security advisories."""

    DEFAULT_PRIMARY_COLOR = "0D6EFD"
    DEFAULT_SECONDARY_COLOR = "6C757D"

    @staticmethod
    def generate(
        threat,
        executive_summary,
        business_impact,
        it_recommendation,
    ):
        settings = WordExporter._load_settings()

        company_name = settings.get(
            "CompanyName",
            "Company Name",
        )
        company_short_name = settings.get(
            "CompanyShortName",
            "",
        )
        department = settings.get(
            "Department",
            "IT Department",
        )
        classification = settings.get(
            "Classification",
            "Internal Use",
        )
        paper_size = settings.get(
            "PaperSize",
            "A4",
        )
        company_logo = settings.get(
            "CompanyLogo",
            "",
        )

        primary_color = WordExporter._parse_hex_color(
            settings.get("PrimaryColor"),
            WordExporter.DEFAULT_PRIMARY_COLOR,
        )
        secondary_color = WordExporter._parse_hex_color(
            settings.get("SecondaryColor"),
            WordExporter.DEFAULT_SECONDARY_COLOR,
        )

        document = Document()

        WordExporter._configure_document(
            document=document,
            paper_size=paper_size,
        )

        WordExporter._configure_styles(
            document=document,
            primary_color=primary_color,
        )

        WordExporter._build_header(
            document=document,
            company_name=company_name,
            company_short_name=company_short_name,
            department=department,
            company_logo=company_logo,
            primary_color=primary_color,
        )

        WordExporter._build_footer(
            document=document,
            classification=classification,
            company_name=company_name,
            secondary_color=secondary_color,
        )

        WordExporter._add_classification_banner(
            document=document,
            classification=classification,
            primary_color=primary_color,
        )

        WordExporter._add_title(
            document=document,
            threat=threat,
            primary_color=primary_color,
        )

        WordExporter._add_threat_information(
            document=document,
            threat=threat,
            primary_color=primary_color,
        )

        WordExporter._add_content_section(
            document=document,
            heading="Executive Summary",
            content=(
                executive_summary
                or "No executive summary available."
            ),
            primary_color=primary_color,
        )

        WordExporter._add_business_impact(
            document=document,
            business_impact=business_impact,
            primary_color=primary_color,
        )

        WordExporter._add_content_section(
            document=document,
            heading="IT Recommendation",
            content=(
                it_recommendation
                or "No IT recommendation available."
            ),
            primary_color=primary_color,
        )

        WordExporter._add_source_reference(
            document=document,
            threat=threat,
            primary_color=primary_color,
        )

        WordExporter._set_document_properties(
            document=document,
            company_name=company_name,
            department=department,
            threat=threat,
        )

        output = BytesIO()
        document.save(output)
        output.seek(0)

        cve = threat.get("CVE") or "Unknown"
        safe_cve = WordExporter._safe_filename(str(cve))

        filename = f"Security_Advisory_{safe_cve}.docx"

        return output, filename

    @staticmethod
    def _load_settings():
        settings = (
            SystemSetting.query
            .filter_by(IsActive=True)
            .all()
        )

        return {
            setting.SettingKey: setting.SettingValue or ""
            for setting in settings
        }

    @staticmethod
    def _configure_document(document, paper_size):
        section = document.sections[0]

        if str(paper_size).upper() == "LETTER":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        else:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)

        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

        section.start_type = WD_SECTION.NEW_PAGE

    @staticmethod
    def _configure_styles(document, primary_color):
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10)

        normal_style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial",
        )

        for style_name, size in [
            ("Title", 22),
            ("Heading 1", 16),
            ("Heading 2", 13),
        ]:
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.color.rgb = primary_color
            style.font.bold = True

            style._element.rPr.rFonts.set(
                qn("w:eastAsia"),
                "Arial",
            )

    @staticmethod
    def _build_header(
        document,
        company_name,
        company_short_name,
        department,
        company_logo,
        primary_color,
    ):
        section = document.sections[0]
        header = section.header

        table = header.add_table(
            rows=1,
            cols=2,
            width=section.page_width,
        )
        table.autofit = False

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        left_cell.width = Cm(4.0)
        right_cell.width = Cm(12.5)

        left_cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )
        right_cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        logo_path = WordExporter._resolve_logo_path(company_logo)

        if logo_path:
            logo_paragraph = left_cell.paragraphs[0]
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = logo_paragraph.add_run()

            try:
                run.add_picture(
                    str(logo_path),
                    height=Cm(1.4),
                )
            except (OSError, ValueError):
                WordExporter._add_company_short_name(
                    paragraph=logo_paragraph,
                    company_short_name=company_short_name,
                    primary_color=primary_color,
                )
        else:
            logo_paragraph = left_cell.paragraphs[0]

            WordExporter._add_company_short_name(
                paragraph=logo_paragraph,
                company_short_name=company_short_name,
                primary_color=primary_color,
            )

        company_paragraph = right_cell.paragraphs[0]
        company_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        company_paragraph.paragraph_format.space_after = Pt(0)

        company_run = company_paragraph.add_run(
            company_name or "Company Name"
        )
        company_run.bold = True
        company_run.font.name = "Arial"
        company_run.font.size = Pt(11)
        company_run.font.color.rgb = primary_color

        department_paragraph = right_cell.add_paragraph()
        department_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        department_paragraph.paragraph_format.space_before = Pt(0)
        department_paragraph.paragraph_format.space_after = Pt(0)

        department_run = department_paragraph.add_run(
            department or "IT Department"
        )
        department_run.font.name = "Arial"
        department_run.font.size = Pt(9)

        WordExporter._remove_table_borders(table)

        separator = header.add_paragraph()
        separator.paragraph_format.space_before = Pt(2)
        separator.paragraph_format.space_after = Pt(0)

        WordExporter._set_paragraph_bottom_border(
            paragraph=separator,
            color=WordExporter._rgb_to_hex(primary_color),
            size="16",
        )

    @staticmethod
    def _build_footer(
        document,
        classification,
        company_name,
        secondary_color,
    ):
        section = document.sections[0]
        footer = section.footer

        separator = footer.paragraphs[0]
        separator.paragraph_format.space_before = Pt(0)
        separator.paragraph_format.space_after = Pt(2)

        WordExporter._set_paragraph_top_border(
            paragraph=separator,
            color=WordExporter._rgb_to_hex(secondary_color),
            size="8",
        )

        table = footer.add_table(
            rows=1,
            cols=3,
            width=section.page_width,
        )
        table.autofit = False

        left_cell = table.cell(0, 0)
        center_cell = table.cell(0, 1)
        right_cell = table.cell(0, 2)

        left_cell.width = Cm(6)
        center_cell.width = Cm(6)
        right_cell.width = Cm(4)

        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        left_run = left_paragraph.add_run(
            company_name or "Company Name"
        )
        WordExporter._format_footer_run(
            run=left_run,
            color=secondary_color,
        )

        center_paragraph = center_cell.paragraphs[0]
        center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        center_run = center_paragraph.add_run(
            classification or "Internal Use"
        )
        center_run.bold = True

        WordExporter._format_footer_run(
            run=center_run,
            color=secondary_color,
        )

        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        page_label_run = right_paragraph.add_run("Page ")
        WordExporter._format_footer_run(
            run=page_label_run,
            color=secondary_color,
        )

        WordExporter._add_page_number_field(right_paragraph)

        WordExporter._remove_table_borders(table)

        generated_paragraph = footer.add_paragraph()
        generated_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_paragraph.paragraph_format.space_before = Pt(0)
        generated_paragraph.paragraph_format.space_after = Pt(0)

        generated_run = generated_paragraph.add_run(
            "Generated by CTI Platform"
        )
        WordExporter._format_footer_run(
            run=generated_run,
            color=secondary_color,
        )

    @staticmethod
    def _add_classification_banner(
        document,
        classification,
        primary_color,
    ):
        table = document.add_table(
            rows=1,
            cols=1,
        )
        table.autofit = True

        cell = table.cell(0, 0)
        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        WordExporter._set_cell_shading(
            cell=cell,
            fill=WordExporter._rgb_to_hex(primary_color),
        )

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)

        run = paragraph.add_run(
            str(classification or "Internal Use").upper()
        )
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

        document.add_paragraph().paragraph_format.space_after = Pt(0)

    @staticmethod
    def _add_title(
        document,
        threat,
        primary_color,
    ):
        title = document.add_paragraph()
        title.style = document.styles["Title"]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)

        title_run = title.add_run(
            "Cyber Security Advisory"
        )
        title_run.font.color.rgb = primary_color

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Pt(0)
        subtitle.paragraph_format.space_after = Pt(14)

        cve = threat.get("CVE") or "No CVE assigned"
        severity = threat.get("Severity") or "Unknown"

        subtitle_run = subtitle.add_run(
            f"{cve} | Severity: {severity}"
        )
        subtitle_run.bold = True
        subtitle_run.font.name = "Arial"
        subtitle_run.font.size = Pt(11)

    @staticmethod
    def _add_threat_information(
        document,
        threat,
        primary_color,
    ):
        WordExporter._add_section_heading(
            document=document,
            text="Threat Information",
            primary_color=primary_color,
        )

        threat_table = document.add_table(
            rows=0,
            cols=2,
        )
        threat_table.style = "Table Grid"
        threat_table.autofit = False

        threat_information = [
            ("Title", threat.get("Title") or "N/A"),
            ("CVE", threat.get("CVE") or "N/A"),
            ("Severity", threat.get("Severity") or "N/A"),
            ("CVSS", threat.get("CVSS") or "N/A"),
            (
                "Known Exploited Vulnerability",
                "Yes" if threat.get("KEV") else "No",
            ),
            ("Source", threat.get("Source") or "N/A"),
            (
                "Published Date",
                WordExporter._format_date(
                    threat.get("PublishedDate")
                ),
            ),
        ]

        for label, value in threat_information:
            row_cells = threat_table.add_row().cells

            row_cells[0].width = Cm(5)
            row_cells[1].width = Cm(11)

            row_cells[0].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )
            row_cells[1].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            WordExporter._set_cell_shading(
                cell=row_cells[0],
                fill=WordExporter._lighten_color(primary_color),
            )

            label_paragraph = row_cells[0].paragraphs[0]
            label_run = label_paragraph.add_run(str(label))
            label_run.bold = True
            label_run.font.name = "Arial"
            label_run.font.size = Pt(9)

            value_paragraph = row_cells[1].paragraphs[0]
            value_run = value_paragraph.add_run(str(value))
            value_run.font.name = "Arial"
            value_run.font.size = Pt(9)

        document.add_paragraph().paragraph_format.space_after = Pt(0)

    @staticmethod
    def _add_content_section(
        document,
        heading,
        content,
        primary_color,
    ):
        WordExporter._add_section_heading(
            document=document,
            text=heading,
            primary_color=primary_color,
        )

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15

        lines = str(content).splitlines()

        for index, line in enumerate(lines):
            if index > 0:
                paragraph.add_run().add_break()

            run = paragraph.add_run(line)
            run.font.name = "Arial"
            run.font.size = Pt(10)

    @staticmethod
    def _add_business_impact(
        document,
        business_impact,
        primary_color,
    ):
        WordExporter._add_section_heading(
            document=document,
            text="Business Impact",
            primary_color=primary_color,
        )

        if business_impact:
            for item in business_impact:
                paragraph = document.add_paragraph(
                    style="List Bullet",
                )
                paragraph.paragraph_format.space_after = Pt(3)

                run = paragraph.add_run(str(item))
                run.font.name = "Arial"
                run.font.size = Pt(10)
        else:
            document.add_paragraph(
                "No business impact information available."
            )

    @staticmethod
    def _add_source_reference(
        document,
        threat,
        primary_color,
    ):
        WordExporter._add_section_heading(
            document=document,
            text="Source Reference",
            primary_color=primary_color,
        )

        reference_url = (
            threat.get("ReferenceUrl")
            or "No reference URL available."
        )

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)

        if WordExporter._is_valid_web_url(reference_url):
            WordExporter._add_hyperlink(
                paragraph=paragraph,
                text=str(reference_url),
                url=str(reference_url),
                color=WordExporter._rgb_to_hex(primary_color),
            )
        else:
            run = paragraph.add_run(str(reference_url))
            run.font.name = "Arial"
            run.font.size = Pt(9)

        generated_paragraph = document.add_paragraph()
        generated_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        generated_paragraph.paragraph_format.space_before = Pt(12)

        generated_run = generated_paragraph.add_run(
            "Generated: "
            f"{datetime.now().strftime('%d %b %Y %H:%M')}"
        )
        generated_run.italic = True
        generated_run.font.name = "Arial"
        generated_run.font.size = Pt(8)

    @staticmethod
    def _add_section_heading(
        document,
        text,
        primary_color,
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)

        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.color.rgb = primary_color

        WordExporter._set_paragraph_bottom_border(
            paragraph=paragraph,
            color=WordExporter._rgb_to_hex(primary_color),
            size="8",
        )

    @staticmethod
    def _add_company_short_name(
        paragraph,
        company_short_name,
        primary_color,
    ):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run(
            company_short_name or "CTI"
        )
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.font.color.rgb = primary_color

    @staticmethod
    def _resolve_logo_path(company_logo):
        if not company_logo:
            return None

        raw_path = str(company_logo).strip()

        if not raw_path:
            return None

        normalized_path = raw_path.replace("\\", "/")

        if normalized_path.startswith("/static/"):
            relative_path = normalized_path.removeprefix("/static/")
            candidate = (
                Path(current_app.static_folder)
                / relative_path
            )
        elif normalized_path.startswith("static/"):
            relative_path = normalized_path.removeprefix("static/")
            candidate = (
                Path(current_app.static_folder)
                / relative_path
            )
        else:
            candidate = Path(raw_path)

            if not candidate.is_absolute():
                candidate = (
                    Path(current_app.root_path)
                    / candidate
                )

        try:
            resolved_candidate = candidate.resolve()
        except OSError:
            return None

        if resolved_candidate.is_file():
            return resolved_candidate

        return None

    @staticmethod
    def _set_document_properties(
        document,
        company_name,
        department,
        threat,
    ):
        properties = document.core_properties

        properties.title = (
            f"Cyber Security Advisory - "
            f"{threat.get('CVE') or 'Unknown'}"
        )
        properties.subject = (
            threat.get("Title")
            or "Cyber Security Advisory"
        )
        properties.author = department or "IT Department"
        properties.company = company_name or "Company Name"
        properties.comments = "Generated by CTI Platform"
        properties.keywords = (
            "CTI, Cybersecurity, Security Advisory, "
            f"{threat.get('CVE') or ''}"
        )

    @staticmethod
    def _add_page_number_field(paragraph):
        run = paragraph.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(8)

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = "PAGE"

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        value = OxmlElement("w:t")
        value.text = "1"

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run._r.append(begin)
        run._r.append(instruction)
        run._r.append(separate)
        run._r.append(value)
        run._r.append(end)

    @staticmethod
    def _add_hyperlink(
        paragraph,
        text,
        url,
        color,
    ):
        part = paragraph.part

        relationship_id = part.relate_to(
            url,
            (
                "http://schemas.openxmlformats.org/"
                "officeDocument/2006/relationships/hyperlink"
            ),
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        run_color = OxmlElement("w:color")
        run_color.set(qn("w:val"), color)
        run_properties.append(run_color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(underline)

        run_fonts = OxmlElement("w:rFonts")
        run_fonts.set(qn("w:ascii"), "Arial")
        run_fonts.set(qn("w:hAnsi"), "Arial")
        run_properties.append(run_fonts)

        run.append(run_properties)

        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _set_cell_shading(cell, fill):
        cell_properties = cell._tc.get_or_add_tcPr()

        shading = cell_properties.find(
            qn("w:shd")
        )

        if shading is None:
            shading = OxmlElement("w:shd")
            cell_properties.append(shading)

        shading.set(qn("w:fill"), fill)

    @staticmethod
    def _remove_table_borders(table):
        table_properties = table._tbl.tblPr

        borders = table_properties.find(
            qn("w:tblBorders")
        )

        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)

        for edge in [
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ]:
            border = borders.find(qn(f"w:{edge}"))

            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)

            border.set(qn("w:val"), "nil")

    @staticmethod
    def _set_paragraph_bottom_border(
        paragraph,
        color,
        size,
    ):
        paragraph_properties = (
            paragraph._p.get_or_add_pPr()
        )

        borders = paragraph_properties.find(
            qn("w:pBdr")
        )

        if borders is None:
            borders = OxmlElement("w:pBdr")
            paragraph_properties.append(borders)

        bottom = borders.find(qn("w:bottom"))

        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)

        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)

    @staticmethod
    def _set_paragraph_top_border(
        paragraph,
        color,
        size,
    ):
        paragraph_properties = (
            paragraph._p.get_or_add_pPr()
        )

        borders = paragraph_properties.find(
            qn("w:pBdr")
        )

        if borders is None:
            borders = OxmlElement("w:pBdr")
            paragraph_properties.append(borders)

        top = borders.find(qn("w:top"))

        if top is None:
            top = OxmlElement("w:top")
            borders.append(top)

        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), size)
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), color)

    @staticmethod
    def _format_footer_run(run, color):
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = color

    @staticmethod
    def _parse_hex_color(value, default_value):
        raw_value = str(value or "").strip().lstrip("#")

        if len(raw_value) != 6:
            raw_value = default_value

        try:
            red = int(raw_value[0:2], 16)
            green = int(raw_value[2:4], 16)
            blue = int(raw_value[4:6], 16)

            return RGBColor(red, green, blue)
        except (TypeError, ValueError):
            return WordExporter._parse_hex_color(
                default_value,
                WordExporter.DEFAULT_PRIMARY_COLOR,
            )

    @staticmethod
    def _rgb_to_hex(color):
        return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

    @staticmethod
    def _lighten_color(color):
        red = min(255, int(color[0] + (255 - color[0]) * 0.88))
        green = min(
            255,
            int(color[1] + (255 - color[1]) * 0.88),
        )
        blue = min(
            255,
            int(color[2] + (255 - color[2]) * 0.88),
        )

        return f"{red:02X}{green:02X}{blue:02X}"

    @staticmethod
    def _format_date(value):
        if value is None:
            return "N/A"

        if hasattr(value, "strftime"):
            return value.strftime("%d %b %Y")

        return str(value)

    @staticmethod
    def _is_valid_web_url(value):
        try:
            parsed = urlparse(str(value))

            return (
                parsed.scheme in {"http", "https"}
                and bool(parsed.netloc)
            )
        except ValueError:
            return False

    @staticmethod
    def _safe_filename(value):
        safe_value = value

        for character in [
            "/",
            "\\",
            ":",
            "*",
            "?",
            '"',
            "<",
            ">",
            "|",
        ]:
            safe_value = safe_value.replace(character, "-")

        return safe_value.strip() or "Unknown"